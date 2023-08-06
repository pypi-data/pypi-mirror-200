import docx
import lxml
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import  Pt
from docx.shared import RGBColor
from docx.shared import Cm, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION


def title_generator(document, title):
    h = document.add_paragraph(style='New Heading')
    h.add_run(title)


def subtitle_generator(document, title):
    h = document.add_paragraph(style='New SubHeading')
    h.add_run(title)


def paragraph_generator(document, content, first_line_indent=True):
    paragraph = document.add_paragraph(style='New Paragraph')
    run = paragraph.add_run(content)

    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def highline_paragraph_generator(document, content, first_line_indent=True):
    paragraph = document.add_paragraph(style='New Paragraph')

    for sub in content:
        run = paragraph.add_run(str(sub[0]))
        if sub[1] == 1:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def highline_artical_generator(document, content, first_line_indent=True):
    for sub in content:
        paragraph = document.add_paragraph(style='New Paragraph')

        run = paragraph.add_run(str(sub[0]))
        if sub[1] == 1:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        if first_line_indent:
            paragraph.paragraph_format.first_line_indent = Cm(0.74)


def no_head_table_generator(document, title, data=[]):
    """
    document: docx object
    title: string
    data: [[<title>, <content>, <is_heighline>, <is_blod>]]
    """
    h = document.add_paragraph(style='Table Heading')
    h.add_run(title)

    table = document.add_table(rows=0, cols=2, style='Table Grid')
    table.allow_autofit = False
    row = table.rows

    for rec in data:
        row = table.add_row()
        row_cells = row.cells
        row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        row.height = Cm(0.48)
        row_cells[0].text = str(rec[0])
        row_cells[1].text = str(rec[1])
        row_cells[0].width = Inches(3.0)
        row_cells[1].width = Inches(6.0)

        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10.5)
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(10.5)
        if len(rec) > 2 and rec[2] == 1:
            row_cells[1].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

        if len(rec) > 3 and rec[3] == 1:
            row_cells[1].paragraphs[0].runs[0].font.bold = True


def table_generator(document, description, head=[], data=[]):
    if isinstance(description, str):
        paragraph_generator(document, description)
    else:
        highline_paragraph_generator(document, description)

    table = document.add_table(rows=1, cols=len(head), style='Table Grid')
    table.allow_autofit = True
    table.rows[0].height = Cm(0.48)
    hdr_cells = table.rows[0].cells

    for i in range(len(head)):
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells[i].text = head[i]
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)

    for rec in range(len(data)):
        row = table.add_row()
        row.height = Cm(0.48)
        row_cells = row.cells
        for i in range(len(data[rec])):
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row_cells[i].text = str(data[rec][i])
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)
            row_cells[i].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW


def crf_table_generator(document, title, data=[]):
    """
    document: docx object
    title: string
    data: [[<title>, <content>, <is_merged_cells>, <is_blod>]]
    """
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading(title)
    section = new_section #document.sections[len(document.sections)-1]

    header = section.header
    footer = section.footer
    header.is_linked_to_previous = False
    footer.is_linked_to_previous = False
    section.different_first_page_header_footer = False
    add_page_number(footer.paragraphs[0].add_run())

    header.paragraphs[0].add_run(f"\t\t{title}")
    header.paragraphs[0].style = document.styles["Header"]
    header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table = document.add_table(rows=0, cols=2, style='Table Grid')
    table.allow_autofit = False
    row = table.rows

    for rec in data:
        row = table.add_row()
        row_cells = row.cells
        row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        row.height = Cm(0.48)
        row_cells[0].text = str(rec[0])
        row_cells[1].text = str(rec[1])
        row_cells[0].width = Inches(3.0)
        row_cells[1].width = Inches(6.0)

        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10.5)
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(10.5)
        if len(rec) > 2 and rec[2] == 1:
            a, b = row_cells[:2]
            a.merge(b)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            if len(rec) > 3 and rec[3] == 1:
                row_cells[0].paragraphs[0].runs[0].font.bold = True


def cover_generator(document, title, data=[]):
    CStyle = document.styles['Cover Heading']
    h = document.add_paragraph(title)
    h.style = CStyle
    # add the table
    table = document.add_table(rows=0, cols=2, style='Table Grid')
    for title, value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(title)
        row_cells[1].text = str(value)
        row_cells[0].width = Inches(2.0)
        row_cells[1].width = Inches(3.0)
        row_cells[0].paragraphs[0].runs[0].font.bold = True

    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_page_break()


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def create_doc():
    document = Document()

    # 正文
    document.styles["Normal"].font.name = u"Times New Roman"  # 设置全局字体
    document.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles["Normal"].font.size = Pt(12)
    document.styles["Normal"].paragraph_format.space_before = Pt(0.5)
    document.styles["Normal"].paragraph_format.space_after = Pt(0.5)

    styles = document.styles

    # paragraph [New Paragraph]
    new_paragraph_style = styles.add_style('New Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    new_paragraph_style.base_style = styles['Normal']
    new_paragraph_style.paragraph_format.alignment = 3
    new_paragraph_style.paragraph_format.line_spacing = 1
    new_paragraph_style.paragraph_format.space_before = Pt(0.5)
    new_paragraph_style.paragraph_format.space_after = Pt(0.5)

    # 一级标题 【New Heading】
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 2']
    new_heading_style.font.name = u"Times New Roman"
    new_heading_style.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色
    new_heading_style.font.size = Pt(14)
    new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    new_heading_style.paragraph_format.first_line_indent = 0
    new_heading_style.paragraph_format.space_before = Pt(10)
    new_heading_style.paragraph_format.space_after = Pt(5)

    # 二级标题 【New SubHeading】
    new_heading_style = styles.add_style('New SubHeading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 2']
    new_heading_style.font.name = u"Times New Roman"
    new_heading_style.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色
    new_heading_style.font.size = Pt(12)
    new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    new_heading_style.paragraph_format.first_line_indent = 0
    new_heading_style.paragraph_format.space_before = Pt(5)
    new_heading_style.paragraph_format.space_after = Pt(0.5)

    # 表格标题【Table Heading】
    new_heading_style = styles.add_style('Table Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 2']
    new_heading_style.font.name = u"Times New Roman"
    new_heading_style.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色
    new_heading_style.font.size = Pt(12)
    new_heading_style.paragraph_format.alignment = 1
    new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    new_heading_style.paragraph_format.space_before = Pt(10)
    new_heading_style.paragraph_format.space_after = Pt(0.5)
    new_heading_style.paragraph_format.first_line_indent = 0

    # Cover 标题【Cover Heading】
    new_heading_style = styles.add_style('Cover Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Normal']
    new_heading_style.font.name = u"Times New Roman"
    new_heading_style.font.color.rgb = RGBColor(42, 75, 113)  # 设置颜色为黑色
    new_heading_style.font.size = Pt(58)
    new_heading_style.paragraph_format.alignment = 1
    new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    new_heading_style.paragraph_format.space_before = Pt(100)
    new_heading_style.paragraph_format.space_after = Pt(20)
    new_heading_style.paragraph_format.first_line_indent = 0

    # ToC 标题【ToC Heading】
    new_heading_style = styles.add_style('ToC Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Normal']
    new_heading_style.font.name = u"Times New Roman"
    new_heading_style.font.color.rgb = RGBColor(42, 75, 113)  # 设置颜色为黑色
    new_heading_style.font.size = Pt(18)
    new_heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    new_heading_style.paragraph_format.space_before = Pt(20)
    new_heading_style.paragraph_format.space_after = Pt(20)
    new_heading_style.paragraph_format.first_line_indent = 0
    return document


def place_table_of_contents(document):

    # Code for making Table of Contents
    paragraph = document.add_paragraph('Table of Contents')
    run = paragraph.add_run()
    paragraph.style = document.styles['ToC Heading']
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:val')
    fldChar3.text = "Right-click to update field."

    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p




